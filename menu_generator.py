"""
Weekly menu generator (Viikkomenu)
Produces a Word document in the restaurant's lunch-list layout:

    Ruokalista
    Lounaslista 24/26

    Ma 8.6    SOUP (L, GL)
              SALAD & SALAATTIBUFFET
              MAIN (L, GL)
              SIDES

    ... Ma-Pe ...

    [footer: contact, hours, prices, diet legend, disclaimers]

The header/footer texts live in MENU_CONFIG below — edit them once
to match the restaurant, then every export uses them.
"""

import re
from datetime import date, timedelta

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ----------------------------------------------------------------------
# EDIT THIS BLOCK to match the restaurant (taken from the Kesti template)
# ----------------------------------------------------------------------
MENU_CONFIG = {
    'title': 'Ruokalista',
    'subtitle': 'Lounaslista {week}/{yy}',
    'salad_line': 'PÄIVÄN SALAATTI & SALAATTIBUFFET',
    'sides_line': 'LISÄKKEET: ',          # left open for the manager to fill
    'contact_lines': [
        'Siltakatu 11, 18100 Heinola',
        'Puh. 050 405 5803 | keittiö',
        'Jyrankola.fi',
    ],
    'hours_lines': [
        'Avoinna ma–pe klo 08.00–14.30',
        'Lounas klo 11–14',
    ],
    'price_lines': [
        'Lounas\t13,70 €',
        'Lounas, eläkeläiset\t12,50 €',
        'Keittolounas\t11,50 €',
        'Salaattilounas\t10,50 €',
        'Lounas, lapset 6–12 v\t8,50 €',
        'Lounas, lapset 1–5 v\t1 €/ikävuosi',
    ],
    'legend': 'L = laktoositon,  GL = gluteeniton,  VEG = vegaaninen,  M = maidoton',
    'disclaimers': [
        'Huom.! Pidämme oikeuden ruokalistan muutoksiin. Tarkemmat tiedot elintarvikkeista '
        'ja niiden sisältämistä allergeeneista saa tarvittaessa henkilökunnalta.',
        'Käytämme suomalaista naudan- ja sianlihaa sekä siipikarjanlihaa.',
    ],
}

DAY_NAMES = ['Ma', 'Ti', 'Ke', 'To', 'Pe']

# PoweResta/system codes -> menu display codes (per the template legend)
CODE_MAP = {'G': 'GL', 'VG': 'VEG'}
MENU_CODES = ['L', 'GL', 'M', 'VEG']   # codes shown on the menu, in this order


def extract_menu_codes(notes):
    """'Ruokavaliot: PERUS,M,L,G,...' -> '(L, GL, M)' or '' if unknown."""
    if not notes:
        return ''
    m = re.search(r'Ruokavaliot:\s*([A-ZÄÖÅ,\-\s]+)', notes)
    if not m:
        return ''
    raw = [c.strip().upper() for c in m.group(1).split(',') if c.strip()]
    mapped = [CODE_MAP.get(c, c) for c in raw]
    shown = [c for c in MENU_CODES if c in mapped]
    return f"({', '.join(shown)})" if shown else ''


def monday_of_week(week_number, year):
    """ISO week -> Monday's date."""
    try:
        return date.fromisocalendar(year, min(week_number, 52), 1)
    except ValueError:
        return date.fromisocalendar(year, 1, 1)


def build_week_menu(meals, week_number, year, output_path, day_names=None):
    """
    meals: EITHER a flat list of dicts {'name':..., 'notes':..., 'category':...}
           (heuristic split: soups = name contains 'keitto', rest = mains;
           no real salad data, generic salad_line used)
       OR   a list of day dicts {'soup','main','salad'} (each a meal dict or
           None), one per weekday in order — the exact, real per-day
           assignment, used as-is.
    day_names: weekday labels, one per day of `meals` — defaults to the
        5-day Ma-Pe list; pass a 7-item list (Ma..Su) for facilities that
        run Monday-Sunday (hoiva, kymenkartano).
    Returns (output_path, warnings).
    """
    warnings = []
    names = day_names or DAY_NAMES
    if meals and isinstance(meals[0], dict) and 'soup' in meals[0]:
        days_data = meals
    else:
        n = len(names)
        soups = [m for m in meals if 'keitto' in m['name'].lower()]
        mains = [m for m in meals if 'keitto' not in m['name'].lower()]
        if len(soups) < n:
            warnings.append(f'Viikolla vain {len(soups)} keittoa — päivät täytetään kiertäen.')
        if len(mains) < n:
            warnings.append(f'Viikolla vain {len(mains)} pääruokaa — päivät täytetään kiertäen.')
        if not soups:
            soups = [{'name': '(KEITTO PUUTTUU)', 'notes': ''}]
        if not mains:
            mains = [{'name': '(PÄÄRUOKA PUUTTUU)', 'notes': ''}]
        days_data = [{'soup': soups[i % len(soups)], 'main': mains[i % len(mains)], 'salad': None}
                     for i in range(n)]

    monday = monday_of_week(week_number, year)

    doc = Document()

    # Page setup: A4, modest margins
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    for side in ('top_margin', 'bottom_margin', 'left_margin', 'right_margin'):
        setattr(section, side, Cm(1.5))

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ---- Header ----
    p = doc.add_paragraph()
    r = p.add_run(MENU_CONFIG['title'])
    r.font.size = Pt(36)
    r.font.bold = True

    p = doc.add_paragraph()
    r = p.add_run(MENU_CONFIG['subtitle'].format(week=week_number, yy=str(year)[-2:]))
    r.font.size = Pt(20)
    r.font.bold = True
    p.space_after = Pt(6)

    # ---- Day blocks: 2-column table (day | dishes), borderless ----
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(3.2)
    table.columns[1].width = Cm(14.8)

    for i, day in enumerate(names):
        d = monday + timedelta(days=i)
        day_info = days_data[i] if i < len(days_data) else {}
        soup = day_info.get('soup')
        main = day_info.get('main')
        salad = day_info.get('salad')

        row = table.add_row()
        day_cell, dish_cell = row.cells
        day_cell.width = Cm(3.2)
        dish_cell.width = Cm(14.8)
        day_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        dp = day_cell.paragraphs[0]
        dr = dp.add_run(f'{day} {d.day}.{d.month}')
        dr.font.size = Pt(16)
        dr.font.bold = True

        soup_line = (f"{soup['name'].upper()} {extract_menu_codes(soup.get('notes'))}".strip()
                     if soup else '(KEITTO PUUTTUU)')
        main_line = (f"{main['name'].upper()} {extract_menu_codes(main.get('notes'))}".strip()
                     if main else '(PÄÄRUOKA PUUTTUU)')
        salad_line = (f"{salad['name'].upper()} & SALAATTIBUFFET" if salad
                     else MENU_CONFIG['salad_line'])

        lines = [soup_line, salad_line, main_line, MENU_CONFIG['sides_line']]
        first = True
        for line in lines:
            lp = dish_cell.paragraphs[0] if first else dish_cell.add_paragraph()
            first = False
            lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            lr = lp.add_run(line)
            lr.font.size = Pt(11)
            lp.paragraph_format.space_after = Pt(2)

        # spacer row between days
        spacer = table.add_row()
        for c in spacer.cells:
            c.paragraphs[0].paragraph_format.space_after = Pt(4)

    # ---- Footer block ----
    doc.add_paragraph()
    foot = doc.add_table(rows=1, cols=3)
    foot.autofit = False
    widths = (Cm(6.0), Cm(5.5), Cm(6.5))
    blocks = (MENU_CONFIG['contact_lines'], MENU_CONFIG['hours_lines'],
              MENU_CONFIG['price_lines'])
    for cell, width, lines in zip(foot.rows[0].cells, widths, blocks):
        cell.width = width
        first = True
        for line in lines:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            r = p.add_run(line)
            r.font.size = Pt(9)

    p = doc.add_paragraph()
    r = p.add_run(MENU_CONFIG['legend'])
    r.font.size = Pt(8)
    r.font.bold = True
    for text in MENU_CONFIG['disclaimers']:
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.save(output_path)
    return output_path, warnings
